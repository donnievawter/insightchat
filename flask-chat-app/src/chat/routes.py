from flask import Blueprint, render_template, request, session, redirect, jsonify
import os
import requests
import asyncio
import os
import requests
from .utils import prompt_model, fetch_repo_chunks, get_available_models, fetch_document_content
from .config import DEFAULT_SYSTEM_PROMPT, DEFAULT_MODEL, DEFAULT_RAG_CHUNKS, MAX_MESSAGE_HISTORY, CSV_ANALYSIS_INSTRUCTIONS, TOOL_SYSTEM_ENABLED
from .whisper_client import WhisperClient
from .tool_router import get_tool_router

chat_bp = Blueprint('chat', __name__)
whisper_client = WhisperClient()

@chat_bp.route("/chat", methods=["GET", "POST"])
def chat():
    if request.method == "GET":
        # Initialize session if needed
        if "message_history" not in session:
            session["message_history"] = []
        system_prompt = request.args.get("context", DEFAULT_SYSTEM_PROMPT)
        session["system_prompt"] = system_prompt
        
        # Get available models from Ollama
        available_models = get_available_models()
        
        # Determine default model from environment or fallback
        default_model = os.getenv("DEFAULT_MODEL", DEFAULT_MODEL)
        
        # If the default model from env is not available, use first available or fallback
        if available_models:
            available_model_names = [model["name"] for model in available_models]
            if default_model not in available_model_names:
                print(f"DEBUG: DEFAULT_MODEL '{default_model}' not found in available models, using first available")
                default_model = available_models[0]["name"]
        
        current_model = session.get("model", default_model)
        current_use_repo_docs = session.get("use_repo_docs", False)
        
        return render_template("chat.html", 
                             message_history=session["message_history"],
                             available_models=available_models,
                             model=current_model,
                             use_repo_docs=current_use_repo_docs,
                             RAG_API_URL=os.getenv("RAG_API_URL", ""))

    if request.method == "POST":
        # Get form data
        model = request.form.get("model", "llama3.2:latest")
        prompt = request.form.get("prompt", "").strip()
        use_repo_docs = bool(request.form.get("use_repo_docs"))
        
        # Check if there's pre-loaded context from Load Source button
        loaded_context = request.form.get("loaded_context")
        loaded_source_meta = request.form.get("loaded_source_meta")
        
        if not prompt:
            return render_template("chat.html", 
                                 message_history=session.get("message_history", []),
                                 error="Please enter a message",
                                 available_models=get_available_models(),
                                 model=model,
                                 use_repo_docs=use_repo_docs,
                                 RAG_API_URL=os.getenv("RAG_API_URL", ""))

        # Initialize message history if not exists
        if "message_history" not in session:
            session["message_history"] = []

        # Store model and RAG preference in session
        session["model"] = model
        session["use_repo_docs"] = use_repo_docs

        # Add user message to history
        session["message_history"].append({"role": "user", "content": prompt})

        # Clean up session to prevent cookie size issues - be more aggressive
        def cleanup_message_history():
            """Keep only the last 6 messages and aggressively remove metadata to stay under 4KB cookie limit"""
            print(f"DEBUG: Session cleanup - current history length: {len(session['message_history'])}")
            
            # More aggressive message limit
            if len(session["message_history"]) > MAX_MESSAGE_HISTORY:
                session["message_history"] = session["message_history"][-MAX_MESSAGE_HISTORY:]
                print(f"DEBUG: Trimmed message history to last {MAX_MESSAGE_HISTORY} messages")
            
            # Remove ALL large metadata from ALL messages except the very last one
            for i, msg in enumerate(session["message_history"]):
                if i < len(session["message_history"]) - 1:  # Keep metadata only for the last message
                    if "rag_chunks" in msg:
                        del msg["rag_chunks"]
                        print(f"DEBUG: Removed rag_chunks from message {i}")
                    if "sources" in msg:
                        del msg["sources"] 
                        print(f"DEBUG: Removed sources from message {i}")
                        
            # Calculate approximate session size
            import json
            import sys
            session_str = json.dumps(dict(session), default=str)
            session_size = sys.getsizeof(session_str.encode('utf-8'))
            print(f"DEBUG: Estimated session size after cleanup: {session_size} bytes")
        
        cleanup_message_history()

        # ========================================================================
        # TOOL ROUTING - Check if external tools should handle this query
        # ========================================================================
        tool_context = ""
        tool_results = []
        
        if TOOL_SYSTEM_ENABLED:
            try:
                print(f"DEBUG: Checking tools for query: {prompt[:100]}")
                router = get_tool_router()
                
                # Route query to appropriate tools (async)
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                tool_results, tool_context = loop.run_until_complete(router.route_query(prompt))
                loop.close()
                
                if tool_results:
                    success_count = sum(1 for r in tool_results if r.get('success'))
                    print(f"DEBUG: Tools executed: {len(tool_results)}, successful: {success_count}")
                    if tool_context:
                        print(f"DEBUG: Tool context length: {len(tool_context)} chars")
                else:
                    print(f"DEBUG: No tools matched this query")
                    
            except Exception as e:
                print(f"DEBUG: Tool routing error: {e}")
                import traceback
                traceback.print_exc()
                # Continue with normal flow even if tools fail
                tool_context = ""
                tool_results = []
        
        # Check if we can reuse context from recent conversation
        recent_full_context = None
        recent_analyzed_docs = []
        
        # Look for recent messages with full document analysis
        for msg in reversed(session["message_history"][-3:]):  # Check last 3 messages
            if msg.get("hybrid_analysis") and msg.get("analyzed_documents"):
                recent_analyzed_docs = msg["analyzed_documents"]
                print(f"DEBUG: Found recent full document analysis: {recent_analyzed_docs}")
                # Check if current query might be about the same documents
                query_lower = prompt.lower()
                doc_keywords = ["document", "portfolio", "csv", "file", "data", "rows", "columns"]
                if any(keyword in query_lower for keyword in doc_keywords):
                    print(f"DEBUG: Follow-up query detected, will reuse recent context")
                    # We'll fetch the same documents again but skip RAG search
                    break
        
        # Fetch context from RAG if requested (but skip if tools already handled the query)
        context_text = None
        rag_chunks = []
        sources_found = []  # Initialize sources list
        
        # Skip RAG if tools successfully handled the query
        tools_handled_query = bool(tool_results and any(r.get('success') for r in tool_results))
        
        if use_repo_docs and not tools_handled_query:
            k = DEFAULT_RAG_CHUNKS  # Default number of chunks
            rag_api_url = os.getenv("RAG_API_URL")
            print(f"DEBUG: RAG enabled, API URL: {rag_api_url}")  # Debug log
            
            if not rag_api_url:
                print("DEBUG: RAG_API_URL not set in environment variables")
                session["message_history"].append({
                    "role": "assistant", 
                    "content": "⚠️ RAG is enabled but RAG_API_URL environment variable is not set. Please configure it in your .env file."
                })
                session.modified = True
                return render_template("chat.html", 
                                     message_history=session["message_history"],
                                     available_models=get_available_models(),
                                     model=model,
                                     use_repo_docs=use_repo_docs,
                                     RAG_API_URL=os.getenv("RAG_API_URL", ""))
            
            # Check if user has pre-loaded context from Load Source button
            if loaded_context:
                print(f"DEBUG: Using pre-loaded context from Load Source button")
                context_text = loaded_context
                rag_chunks = []  # No need for regular RAG chunks
                
                # Parse source metadata if available
                if loaded_source_meta:
                    import json
                    try:
                        meta = json.loads(loaded_source_meta)
                        source_path = meta.get('source_path', '')
                        context_type = meta.get('context_type', 'unknown')
                        print(f"DEBUG: Loaded context from {source_path} ({context_type})")
                        
                        # Add source info for display
                        if source_path:
                            file_ext = source_path.split('.')[-1].lower() if '.' in source_path else ''
                            is_csv = file_ext == 'csv' or context_type == 'csv_full'
                            sources_found.append({
                                'path': source_path,
                                'filename': source_path.split('/')[-1],
                                'is_csv': is_csv,
                                'file_type': file_ext
                            })
                    except json.JSONDecodeError:
                        print("DEBUG: Could not parse loaded source metadata")
            # Decide whether to do new RAG search or reuse recent context
            elif recent_analyzed_docs:
                print(f"DEBUG: Reusing recent document context instead of new RAG search")
                # Skip RAG search, we'll fetch the same documents directly
                analyze_documents = recent_analyzed_docs
                rag_chunks = []  # No new chunks needed
                context_text = None  # No new RAG context needed
            else:
                # Get both context and chunk data via normal RAG search
                context_text, rag_chunks = fetch_repo_chunks(prompt, k=k, rag_api_url=rag_api_url, return_chunks=True)
                print(f"DEBUG: RAG context retrieved: {bool(context_text)}")  # Debug log
                print(f"DEBUG: RAG chunks retrieved: {len(rag_chunks)}")  # Debug log
            
            # Collect source information for later loading if user requests it
            if rag_chunks:
                unique_sources = set()
                for chunk in rag_chunks:
                    source = chunk.get('metadata', {}).get('source', '')
                    if source and source not in unique_sources:
                        unique_sources.add(source)
                        # Determine file type for special handling hints
                        file_ext = source.split('.')[-1].lower() if '.' in source else ''
                        is_csv = file_ext == 'csv'
                        sources_found.append({
                            'path': source,
                            'filename': source.split('/')[-1],
                            'is_csv': is_csv,
                            'file_type': file_ext
                        })
                print(f"DEBUG: Found {len(sources_found)} unique sources in chunks")
            
            # Use the RAG context if we have it
            if context_text:
                combined_context = context_text
                print(f"DEBUG: Using RAG chunks context length: {len(combined_context)} chars")
                
                # Combine tool context with RAG context if both exist
                if tool_context:
                    combined_context = tool_context + "\n\n" + combined_context
                    print(f"DEBUG: Combined tool + RAG context length: {len(combined_context)} chars")
                
                # Insert context as system message temporarily for this query
                temp_history = [{"role": "system", "content": combined_context}] + session["message_history"]
            elif tool_context:
                # Only tool context, no RAG context
                print(f"DEBUG: Using tool context only: {len(tool_context)} chars")
                temp_history = [{"role": "system", "content": tool_context}] + session["message_history"]
            else:
                print("DEBUG: No context retrieved from RAG or tools")
                temp_history = session["message_history"]
        else:
            # RAG not enabled - but check if we have pre-loaded context from Load Source button
            if loaded_context:
                print(f"DEBUG: Using pre-loaded context (RAG disabled but context loaded manually)")
                combined_context = loaded_context
                
                # Parse source metadata if available
                if loaded_source_meta:
                    import json
                    try:
                        meta = json.loads(loaded_source_meta)
                        source_path = meta.get('source_path', '')
                        context_type = meta.get('context_type', 'unknown')
                        print(f"DEBUG: Loaded context from {source_path} ({context_type})")
                        
                        # Add source info for display
                        if source_path:
                            file_ext = source_path.split('.')[-1].lower() if '.' in source_path else ''
                            is_csv = file_ext == 'csv' or context_type == 'csv_full'
                            sources_found.append({
                                'path': source_path,
                                'filename': source_path.split('/')[-1],
                                'is_csv': is_csv,
                                'file_type': file_ext
                            })
                    except json.JSONDecodeError:
                        print("DEBUG: Could not parse loaded source metadata")
                
                # Combine with tool context if present
                if tool_context:
                    combined_context = tool_context + "\n\n" + combined_context
                
                temp_history = [{"role": "system", "content": combined_context}] + session["message_history"]
            elif tool_context:
                print(f"DEBUG: Using tool context only (RAG disabled): {len(tool_context)} chars")
                temp_history = [{"role": "system", "content": tool_context}] + session["message_history"]
            else:
                print("DEBUG: RAG not enabled and no tool context")
                temp_history = session["message_history"]
            
            # Get available documents for Load button functionality
            rag_api_url = os.getenv("RAG_API_URL")
            if rag_api_url:
                try:
                    # Get available documents list
                    response = requests.get(f"{rag_api_url}/documents")
                    if response.status_code == 200:
                        available_docs = response.json().get('documents', [])
                        for doc in available_docs:
                            # Handle both old format (string) and new format (dict)
                            if isinstance(doc, dict):
                                doc_path = doc.get('source', '')
                                file_type = doc.get('file_type', '')
                                filename = doc_path.split('/')[-1] if doc_path else ''
                            else:
                                # Old format - doc is a string path
                                doc_path = doc
                                filename = doc.split('/')[-1]
                                file_type = doc.split('.')[-1].lower() if '.' in doc else ''
                            
                            is_csv = file_type == 'csv'
                            sources_found.append({
                                'path': doc_path,
                                'filename': filename,
                                'is_csv': is_csv,
                                'file_type': file_type
                            })
                        print(f"DEBUG: Found {len(sources_found)} available documents for loading")
                except requests.RequestException as e:
                    print(f"DEBUG: Could not fetch available documents: {e}")

        # Get response from Ollama
        try:
            # Enhanced system prompt for better structured data analysis
            base_system_prompt = session.get("system_prompt", DEFAULT_SYSTEM_PROMPT)
            
            # Add CSV analysis instructions if we have full document context
            if 'full_document_context' in locals() and full_document_context:
                system_prompt = base_system_prompt + CSV_ANALYSIS_INSTRUCTIONS
            else:
                system_prompt = base_system_prompt
            

            
            response_text, _ = prompt_model(
                model=model, 
                prompt=prompt, 
                history=temp_history[:-1],  # Exclude the current user message since it's added in prompt_model
                system_prompt=system_prompt
            )
            
            # Add assistant response to permanent history with sources for potential loading
            assistant_message = {
                "role": "assistant", 
                "content": response_text
            }
            
            # Include RAG chunks if they were used (for backwards compatibility)
            if rag_chunks:
                assistant_message["rag_chunks"] = rag_chunks
            
            # Include sources information for Load button functionality
            if sources_found:
                assistant_message["sources"] = sources_found
                print(f"DEBUG: Stored {len(sources_found)} sources in assistant message")
            
            # Track which tools were used for this response
            if tool_results:
                successful_tools = [r['metadata']['tool'] for r in tool_results if r.get('success')]
                if successful_tools:
                    assistant_message["tools_used"] = successful_tools
                    print(f"DEBUG: Tools used for this response: {successful_tools}")
                
            session["message_history"].append(assistant_message)
            
            # Inline session cleanup to prevent cookie overflow
            print(f"DEBUG: Post-response cleanup - history length: {len(session['message_history'])}")
            if len(session["message_history"]) > MAX_MESSAGE_HISTORY:
                session["message_history"] = session["message_history"][-MAX_MESSAGE_HISTORY:]
                print(f"DEBUG: Trimmed message history to last {MAX_MESSAGE_HISTORY} messages")
            
            # Remove metadata from older messages
            for i, msg in enumerate(session["message_history"]):
                if i < len(session["message_history"]) - 1:
                    if "rag_chunks" in msg:
                        del msg["rag_chunks"]
                    if "sources" in msg:
                        del msg["sources"]
            
            session.modified = True  # Mark session as modified
            
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            session["message_history"].append({"role": "assistant", "content": error_msg})
            
            # Inline session cleanup for error case
            if len(session["message_history"]) > MAX_MESSAGE_HISTORY:
                session["message_history"] = session["message_history"][-MAX_MESSAGE_HISTORY:]
            
            session.modified = True

        # Get available models for the template
        available_models = get_available_models()
        
        return render_template("chat.html", 
                             message_history=session["message_history"],
                             available_models=available_models,
                             model=model,
                             use_repo_docs=use_repo_docs,
                             RAG_API_URL=os.getenv("RAG_API_URL", ""))

@chat_bp.route("/reset", methods=["POST"])
def reset():
    session.clear()
    return redirect("/chat")

@chat_bp.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint for monitoring services"""
    return {"status": "healthy", "service": "insightchat"}, 200

@chat_bp.route("/tools/status", methods=["GET"])
def tools_status():
    """Get status of all available external tools"""
    if not TOOL_SYSTEM_ENABLED:
        return jsonify({
            "enabled": False,
            "message": "Tool system is disabled"
        }), 200
    
    try:
        router = get_tool_router()
        
        # Get tool information
        tool_info = router.get_tool_info()
        active_tools = router.get_active_tools()
        
        # Perform health checks (async)
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        health_status = loop.run_until_complete(router.health_check_all())
        loop.close()
        
        return jsonify({
            "enabled": True,
            "tools": tool_info,
            "active_tools": active_tools,
            "health": health_status
        }), 200
        
    except Exception as e:
        return jsonify({
            "enabled": True,
            "error": str(e)
        }), 500

@chat_bp.route("/browse_documents", methods=["GET"])
def browse_documents():
    """Get list of all documents available in the RAG system for manual selection"""
    from flask import jsonify
    
    rag_api_url = os.getenv("RAG_API_URL")
    if not rag_api_url:
        return jsonify({"error": "RAG API not configured"}), 503
    
    try:
        # Use the new /documents endpoint
        response = requests.get(
            f"{rag_api_url}/documents",
            timeout=15
        )
        
        if response.status_code != 200:
            return jsonify({"error": f"RAG API error: {response.status_code}"}), 500
        
        data = response.json()
        documents = data.get('documents', [])
        total = data.get('total_documents', len(documents))
        
        # Enhance document list with metadata for UI
        enhanced_docs = []
        for doc in documents:
            doc_path = doc.get('source', '')
            filename = doc_path.split('/')[-1]
            file_type = doc.get('file_type', '')
            chunk_count = doc.get('chunk_count', 0)
            
            # Determine file extension from path if not provided
            if not file_type or file_type == 'unknown':
                file_type = doc_path.split('.')[-1].lower() if '.' in doc_path else 'unknown'
            
            is_csv = file_type == 'csv'
            is_ipynb = file_type == 'ipynb'
            
            enhanced_docs.append({
                'path': doc_path,
                'filename': filename,
                'file_type': file_type,
                'is_csv': is_csv,
                'is_ipynb': is_ipynb,
                'chunk_count': chunk_count
            })
        
        print(f"DEBUG: Found {total} documents from RAG API")
        
        return jsonify({
            'success': True,
            'documents': enhanced_docs,
            'count': total
        })
            
    except Exception as e:
        print(f"DEBUG: Error fetching documents: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to fetch documents: {str(e)}"}), 500

@chat_bp.route("/upload_to_rag", methods=["POST"])
def upload_to_rag():
    """Proxy endpoint to upload files to RAG API (avoids CORS issues)"""
    from flask import jsonify
    
    rag_api_url = os.getenv("RAG_API_URL")
    if not rag_api_url:
        return jsonify({"error": "RAG API not configured"}), 503
    
    # Check if file is in request
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    try:
        # Forward the file to RAG API
        files = {'file': (file.filename, file.stream, file.content_type)}
        
        print(f"DEBUG: Uploading file {file.filename} to RAG API")
        
        response = requests.post(
            f"{rag_api_url}/upload",
            files=files,
            timeout=120  # Longer timeout for file upload
        )
        
        print(f"DEBUG: RAG API upload response: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            return jsonify({
                "success": True,
                "message": f"File '{file.filename}' uploaded successfully",
                "details": result
            })
        else:
            error_text = response.text
            print(f"DEBUG: RAG API upload error: {error_text}")
            return jsonify({
                "error": f"RAG API returned status {response.status_code}",
                "details": error_text
            }), response.status_code
            
    except requests.RequestException as e:
        print(f"DEBUG: Error uploading to RAG API: {e}")
        return jsonify({"error": f"Failed to upload file: {str(e)}"}), 500
    except Exception as e:
        print(f"DEBUG: Unexpected error in upload: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Unexpected error: {str(e)}"}), 500

@chat_bp.route("/load_source", methods=["POST"])
def load_source():
    """Load expanded context for a specific source and re-run the last query"""
    from flask import jsonify
    from .utils import fetch_repo_chunks, fetch_document_content
    
    data = request.get_json()
    if not data:
        return jsonify({"error": "No JSON data provided"}), 400
    
    source_path = data.get("source_path")
    
    if not source_path:
        return jsonify({"error": "Missing source_path"}), 400
    
    # Get RAG API URL
    rag_api_url = os.getenv("RAG_API_URL")
    if not rag_api_url:
        return jsonify({"error": "RAG API not configured"}), 503
    
    try:
        # Check if this is a CSV file - if so, load full document
        is_csv = source_path.endswith('.csv')
        
        if is_csv:
            print(f"DEBUG: Loading full CSV document for: {source_path}")
            # For CSV files, load the complete document
            full_content = fetch_document_content(source_path, rag_api_url)
            if full_content and not isinstance(full_content, bytes):
                enhanced_context = f"""PRIORITY: Complete CSV document for comprehensive analysis:

IMPORTANT INSTRUCTIONS FOR CSV ANALYSIS:
- When counting rows, count each line including header (total rows in file)
- When counting properties/records, count data rows only (exclude header)  
- Examine the complete structure systematically
- For row counts: Count every line break to get total rows
- For data counts: Count entries excluding the header row

---
Full Document: {source_path}
{full_content}
---"""
            else:
                return jsonify({"error": "Could not load CSV content"}), 500
        else:
            print(f"DEBUG: Loading expanded chunks for source: {source_path}")
            # Use the new RAG API endpoint to get all chunks for this document
            import requests
            
            try:
                response = requests.post(
                    f"{rag_api_url}/get_chunks_for_document",
                    json={
                        "source": source_path,
                        "limit": 0  # 0 means get all chunks
                    },
                    timeout=30
                )
                
                if response.status_code == 200:
                    chunks_data = response.json()
                    chunks = chunks_data.get('chunks', [])
                    
                    if chunks:
                        # Combine all chunks for this source
                        chunk_contents = [chunk.get('content', '') for chunk in chunks]
                        enhanced_context = f"""Expanded context from all chunks in {source_path}:

{chr(10).join(chunk_contents)}"""
                    else:
                        return jsonify({"error": f"No chunks found for source: {source_path}"}), 404
                else:
                    print(f"DEBUG: RAG API returned status {response.status_code}: {response.text}")
                    return jsonify({"error": f"RAG API error: {response.status_code}"}), 500
                    
            except requests.RequestException as e:
                print(f"DEBUG: Error calling RAG API: {e}")
                return jsonify({"error": f"Failed to fetch chunks: {str(e)}"}), 500
        
        return jsonify({
            "success": True,
            "enhanced_context": enhanced_context,
            "source_path": source_path,
            "context_type": "csv_full" if is_csv else "expanded_chunks"
        })
        
    except Exception as e:
        print(f"DEBUG: Error in load_source: {e}")
        return jsonify({"error": str(e)}), 500

@chat_bp.route("/test", methods=["GET"])
def test_route():
    """Test route to verify blueprint registration"""
    print("TEST ROUTE HIT!")
    return "Test route working", 200

@chat_bp.route("/document", methods=["GET"])
def get_document():
    """Serve document content for the document viewer"""
    print("="*50)
    print("DOCUMENT ROUTE HIT!")
    print("="*50)
    
    from flask import jsonify
    from .utils import fetch_document_content
    
    source = request.args.get("source")
    format_type = request.args.get("format", "raw")  # 'text' or 'raw'
    print(f"DEBUG: Document route called with source: '{source}', format: '{format_type}'")
    print(f"DEBUG: All request args: {dict(request.args)}")
    
    if not source:
        print("DEBUG: No source parameter provided")
        return jsonify({"error": "No source specified"}), 400
    
    # URL decode the source to handle any double encoding
    import urllib.parse
    decoded_source = urllib.parse.unquote(source)
    print(f"DEBUG: Decoded source: '{decoded_source}'")
    
    rag_api_url = os.getenv("RAG_API_URL")
    print(f"DEBUG: RAG_API_URL: {rag_api_url}")
    
    if not rag_api_url:
        print("DEBUG: RAG_API_URL not configured")
        return jsonify({"error": "RAG API not configured"}), 503
    
    try:
        print(f"DEBUG: Attempting to fetch document content for: {decoded_source}")
        content = fetch_document_content(decoded_source, rag_api_url)
        print(f"DEBUG: fetch_document_content returned: {type(content)} with length {len(content) if content else 0}")
        
        if content is None:
            print(f"DEBUG: Content is None, returning 404 for {decoded_source}")
            return jsonify({"error": "Document not found or not accessible"}), 404
        
        # Determine content type based on file extension
        file_extension = decoded_source.split('.')[-1].lower() if '.' in decoded_source else ''
        
        # Handle DOCX text extraction if format=text is requested
        if format_type == 'text' and file_extension == 'docx':
            print(f"DEBUG: Extracting text from DOCX file")
            try:
                from docx import Document
                from io import BytesIO
                
                # Convert to bytes if needed
                if isinstance(content, str):
                    import base64
                    content = base64.b64decode(content)
                
                # Parse DOCX and extract text
                doc = Document(BytesIO(content))
                text_content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_content.append(row_text)
                
                extracted_text = '\n\n'.join(text_content)
                print(f"DEBUG: Successfully extracted {len(extracted_text)} characters from DOCX")
                return extracted_text, 200, {'Content-Type': 'text/plain; charset=utf-8'}
                
            except Exception as docx_error:
                print(f"DEBUG: DOCX extraction failed: {docx_error}")
                import traceback
                print(f"DEBUG: DOCX extraction traceback: {traceback.format_exc()}")
                return jsonify({"error": f"Failed to extract text from DOCX: {str(docx_error)}"}), 500
        content_type_map = {
            'png': 'image/png',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif': 'image/gif',
            'webp': 'image/webp',
            'svg': 'image/svg+xml',
            'bmp': 'image/bmp',
            'pdf': 'application/pdf',
            'json': 'application/json',
            'xml': 'application/xml',
            'html': 'text/html',
            'css': 'text/css',
            'js': 'application/javascript',
            'eml': 'message/rfc822',
            'emlx': 'message/rfc822',
            'wav': 'audio/wav',
            'mp3': 'audio/mpeg',
            'm4a': 'audio/mp4',
            'flac': 'audio/flac',
            'ogg': 'audio/ogg',
        }
        
        content_type = content_type_map.get(file_extension, 'text/plain; charset=utf-8')
        print(f"DEBUG: Using content type: {content_type} for extension: {file_extension}")
        
        # Handle binary vs text content
        if isinstance(content, bytes):
            print(f"DEBUG: Returning binary content ({len(content)} bytes) for {decoded_source}")
            from flask import Response
            return Response(content, 200, {'Content-Type': content_type})
        else:
            print(f"DEBUG: Returning text content ({len(content)} chars) for {decoded_source}")
            return content, 200, {'Content-Type': content_type}
        
    except Exception as e:
        print(f"DEBUG: Exception in document route for {decoded_source}: {type(e).__name__}: {e}")
        import traceback
        print(f"DEBUG: Full traceback: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500

@chat_bp.route("/render_email", methods=["GET"])
def render_email():
    """Render an email file as formatted HTML for viewing in browser"""
    from flask import jsonify
    import requests
    from email import policy
    from email.parser import BytesParser
    import html
    
    source = request.args.get("source")
    print(f"DEBUG: Email render route called with source: '{source}'")
    
    if not source:
        print("DEBUG: No source parameter provided")
        return jsonify({"error": "No source specified"}), 400
    
    # URL decode the source
    import urllib.parse
    decoded_source = urllib.parse.unquote(source)
    print(f"DEBUG: Decoded source: '{decoded_source}'")
    
    rag_api_url = os.getenv("RAG_API_URL")
    print(f"DEBUG: RAG_API_URL: {rag_api_url}")
    
    if not rag_api_url:
        print("DEBUG: RAG_API_URL not configured")
        return jsonify({"error": "RAG API not configured"}), 503
    
    try:
        # Get the raw email file from RAG API's /document endpoint
        from .utils import fetch_document_content
        email_content = fetch_document_content(decoded_source, rag_api_url)
        print(f"DEBUG: Retrieved email file, size: {len(email_content) if email_content else 0} bytes")
        
        if not email_content:
            print("DEBUG: No email content retrieved")
            return jsonify({"error": "Email file not found"}), 404
        
        # Ensure we have bytes for the email parser
        if isinstance(email_content, str):
            email_content = email_content.encode('utf-8')
        
        # Parse the email
        from io import BytesIO
        email_buffer = BytesIO(email_content)
        
        # Mac .emlx files have a header line with message length, skip it
        first_line = email_buffer.readline()
        # If it looks like a length header (just digits), it's .emlx format
        if first_line.strip().isdigit():
            # Continue reading from current position (after the length line)
            msg = BytesParser(policy=policy.default).parse(email_buffer)
        else:
            # Regular .eml file, rewind and parse from beginning
            email_buffer.seek(0)
            msg = BytesParser(policy=policy.default).parse(email_buffer)
        
        print(f"DEBUG: Email parsed successfully")
        
        # Extract metadata
        subject = html.escape(msg.get('subject', '(No Subject)'))
        from_addr = html.escape(msg.get('from', '(Unknown Sender)'))
        to_addr = html.escape(msg.get('to', '(Unknown Recipient)'))
        date = html.escape(msg.get('date', '(No Date)'))
        
        # Extract body content
        body_text = ""
        body_html = None
        
        if msg.is_multipart():
            # Handle multipart messages
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get('Content-Disposition', ''))
                
                # Skip attachments
                if 'attachment' in content_disposition:
                    continue
                    
                # Get plain text parts
                if content_type == 'text/plain':
                    try:
                        text = part.get_content()
                        if text:
                            body_text += text + "\n"
                    except Exception as e:
                        print(f"DEBUG: Error getting plain text part: {e}")
                        pass
                        
                # Get HTML parts
                elif content_type == 'text/html':
                    try:
                        html_content = part.get_content()
                        body_html = html_content
                    except Exception as e:
                        print(f"DEBUG: Error getting HTML part: {e}")
                        pass
        else:
            # Simple non-multipart message
            content_type = msg.get_content_type()
            if content_type == 'text/plain':
                try:
                    body_text = msg.get_content()
                except Exception as e:
                    print(f"DEBUG: Error getting plain text content: {e}")
                    pass
            elif content_type == 'text/html':
                try:
                    body_html = msg.get_content()
                except Exception as e:
                    print(f"DEBUG: Error getting HTML content: {e}")
                    pass
        
        # Render HTML response
        if body_html:
            # Use the HTML version if available
            email_body = f"""
                <div style="border: 1px solid #ddd; padding: 15px; margin-top: 15px; background: white;">
                    {body_html}
                </div>
            """
        else:
            # Use plain text, convert to HTML with line breaks
            escaped_text = html.escape(body_text.strip())
            formatted_text = escaped_text.replace('\n', '<br>\n')
            email_body = f"""
                <div style="border: 1px solid #ddd; padding: 15px; margin-top: 15px; background: white; white-space: pre-wrap; font-family: monospace;">
                    {formatted_text}
                </div>
            """
        
        # Build email content for embedding in modal (not a full HTML page)
        html_content = f"""
<style>
    .email-container {{
        background: white;
        border: 1px solid #ddd;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }}
    .email-header {{
        background: #f8f9fa;
        padding: 20px;
        border-bottom: 2px solid #e9ecef;
    }}
    .email-subject {{
        font-size: 24px;
        font-weight: 600;
        margin: 0 0 15px 0;
        color: #212529;
    }}
    .email-meta {{
        display: grid;
        gap: 8px;
        font-size: 14px;
        color: #495057;
    }}
    .email-meta-row {{
        display: flex;
    }}
    .email-meta-label {{
        font-weight: 600;
        min-width: 80px;
        color: #6c757d;
    }}
    .email-meta-value {{
        flex: 1;
    }}
    .email-body {{
        padding: 20px;
    }}
    .file-info {{
        background: #e9ecef;
        padding: 10px 20px;
        border-top: 1px solid #ddd;
        font-size: 12px;
        color: #6c757d;
    }}
</style>
<div class="email-container">
    <div class="email-header">
        <h1 class="email-subject">{subject}</h1>
        <div class="email-meta">
            <div class="email-meta-row">
                <span class="email-meta-label">From:</span>
                <span class="email-meta-value">{from_addr}</span>
            </div>
            <div class="email-meta-row">
                <span class="email-meta-label">To:</span>
                <span class="email-meta-value">{to_addr}</span>
            </div>
            <div class="email-meta-row">
                <span class="email-meta-label">Date:</span>
                <span class="email-meta-value">{date}</span>
            </div>
        </div>
    </div>
    <div class="email-body">
        {email_body}
    </div>
    <div class="file-info">
        File: {html.escape(decoded_source)}
    </div>
</div>
        """
        
        print(f"DEBUG: Rendered email content, length: {len(html_content)} chars")
        return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
            
    except requests.RequestException as e:
        print(f"DEBUG: Request exception getting email file: {e}")
        return jsonify({"error": f"Failed to retrieve email file: {str(e)}"}), 503
    except Exception as e:
        print(f"DEBUG: Exception in render_email: {type(e).__name__}: {e}")
        return jsonify({"error": f"Error rendering email: {str(e)}"}), 500


@chat_bp.route("/transcribe", methods=["POST"])
def transcribe_audio():
    """
    Transcribe audio file to text using Whisper.
    """
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({"error": "No audio file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Read audio file
        audio_content = file.read()
        
        # Get optional language parameter
        language = request.form.get('language')
        
        # Run async transcription in sync context
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            transcription_result = loop.run_until_complete(
                whisper_client.transcribe(
                    audio_file=audio_content,
                    filename=file.filename or "audio.webm",
                    language=language,
                )
            )
        finally:
            loop.close()
        
        transcribed_text = transcription_result.get("text", "")
        
        if not transcribed_text:
            return jsonify({"error": "No text was transcribed from the audio file"}), 400
        
        print(f"Transcription successful: {transcribed_text[:100]}...")
        
        return jsonify({
            "text": transcribed_text,
            "language": transcription_result.get("language"),
            "duration": transcription_result.get("duration"),
        })
        
    except Exception as e:
        print(f"Error transcribing audio: {str(e)}")
        import traceback
        print(f"DEBUG: Full traceback: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@chat_bp.route("/api/voice-query", methods=["POST"])
def voice_query():
    """
    Voice assistant API endpoint.
    
    Accepts:
    - Audio file (transcribes via Whisper)
    - OR pre-transcribed text
    
    Processes through full chat pipeline (tools + RAG + LLM)
    
    Optionally broadcasts response to TTS speakers
    
    Request:
        - file: audio file (multipart/form-data)
        - OR text: pre-transcribed text (JSON)
        - model: optional model name (defaults to DEFAULT_MODEL)
        - use_rag: optional boolean (defaults to true)
        - broadcast: optional boolean (if true, sends response to TTS)
        - language: optional language code for transcription
    
    Response:
        {
            "success": true/false,
            "query": "transcribed or provided text",
            "response": "LLM response text",
            "tools_used": ["calendar", "weather"],
            "broadcast_sent": true/false,
            "error": "error message if failed"
        }
    """
    try:
        transcribed_text = None
        
        # Check if audio file was uploaded
        if 'file' in request.files:
            file = request.files['file']
            if file.filename:
                print("DEBUG: Voice query - transcribing audio file")
                # Read audio file
                audio_content = file.read()
                language = request.form.get('language')
                
                # Run async transcription
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    transcription_result = loop.run_until_complete(
                        whisper_client.transcribe(
                            audio_file=audio_content,
                            filename=file.filename or "audio.webm",
                            language=language,
                        )
                    )
                finally:
                    loop.close()
                
                transcribed_text = transcription_result.get("text", "").strip()
                if not transcribed_text:
                    return jsonify({
                        "success": False,
                        "error": "No text was transcribed from the audio file"
                    }), 400
                
                print(f"DEBUG: Transcribed: {transcribed_text[:100]}...")
        
        # If no audio, check for text in JSON body
        if not transcribed_text:
            data = request.get_json() or {}
            transcribed_text = data.get('text', '').strip()
            if not transcribed_text:
                return jsonify({
                    "success": False,
                    "error": "No audio file or text provided"
                }), 400
        
        # Get parameters (from form data or JSON)
        if request.content_type and 'multipart/form-data' in request.content_type:
            model = request.form.get('model', os.getenv("DEFAULT_MODEL", DEFAULT_MODEL))
            use_rag = request.form.get('use_rag', 'true').lower() == 'true'
            broadcast = request.form.get('broadcast', 'false').lower() == 'true'
            tts_speaker = request.form.get('speaker')
            tts_model = request.form.get('tts_model')
            tts_engine = request.form.get('engine')
        else:
            data = request.get_json() or {}
            model = data.get('model', os.getenv("DEFAULT_MODEL", DEFAULT_MODEL))
            use_rag = data.get('use_rag', True)
            broadcast = data.get('broadcast', False)
            tts_speaker = data.get('speaker')
            tts_model = data.get('tts_model')
            tts_engine = data.get('engine')
        
        print(f"DEBUG: Voice query - Model: {model}, RAG: {use_rag}, Broadcast: {broadcast}")
        print(f"DEBUG: Query: {transcribed_text}")
        
        # === TOOL ROUTING ===
        tool_context = ""
        tool_results = []
        tools_used = []
        
        if TOOL_SYSTEM_ENABLED:
            router = get_tool_router()
            print(f"DEBUG: Checking tools for query: {transcribed_text}")
            
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                tool_results, tool_context = loop.run_until_complete(
                    router.route_query(transcribed_text)
                )
            finally:
                loop.close()
            
            if tool_results:
                successful_tools = [r['metadata']['tool'] for r in tool_results if r.get('success')]
                tools_used = successful_tools
                print(f"DEBUG: Tools executed: {len(tool_results)}, successful: {len(successful_tools)}")
        
        # === RAG PROCESSING ===
        context_text = None
        tools_handled_query = bool(tool_results and any(r.get('success') for r in tool_results))
        
        if use_rag and not tools_handled_query:
            rag_api_url = os.getenv("RAG_API_URL")
            if rag_api_url:
                print(f"DEBUG: Fetching RAG context for voice query")
                context_text = fetch_repo_chunks(transcribed_text, k=DEFAULT_RAG_CHUNKS, rag_api_url=rag_api_url, return_chunks=False)
                if context_text:
                    print(f"DEBUG: RAG context retrieved: {len(context_text)} chars")
        
        # === PREPARE MESSAGE HISTORY ===
        # Voice queries are stateless - no session history
        message_history = [{
            "role": "user",
            "content": transcribed_text
        }]
        
        # === CALL LLM ===
        print(f"DEBUG: Calling LLM with model: {model}")
        
        # Use a more conversational system prompt for voice
        voice_system_prompt = '''You are a helpful voice assistant. 
        Provide clear, concise responses suitable for speech. 
        Do not include emoticons or characters that cannot be spoken.
        Keep answers brief unless detail is specifically requested. 
        Do not include asterisks or * or ** in your response. 
        Respond in a natural, conversational tone.
        Do not use markdown formatting.
        If the response contains currency amounts, read them out loud with the currency name. For example, "$20" should be read as "20 dollars".'''
        
        # Build context for this query, prepending voice instructions
        combined_context = voice_system_prompt
        
        if tool_context:
            combined_context += "\n\n" + tool_context
        
        if context_text and not tools_handled_query:
            combined_context += "\n\n" + context_text
        
        # Prepare temp history with combined context (voice prompt + tool/RAG data)
        temp_history = [{"role": "system", "content": combined_context}] + message_history
        
        response_text, _ = prompt_model(
            model=model,
            prompt=transcribed_text,
            history=temp_history[:-1],
            system_prompt=""  # Empty string signals system message already in history
        )
        
        print(f"DEBUG: LLM response length: {len(response_text)} chars")
        
        # === OPTIONAL TTS BROADCAST ===
        broadcast_sent = False
        if broadcast:
            tts_url = os.getenv('TTS_BROADCAST_URL')
            if tts_url and tts_speaker:
                try:
                    print(f"DEBUG: Broadcasting response to TTS: {tts_url}")
                    print(f"DEBUG: TTS speaker: {tts_speaker}, model: {tts_model or 'default'}")
                    tts_timeout = int(os.getenv('TTS_TIMEOUT', '10'))
                    
                    # Build TTS request payload
                    # Limit TTS text length to prevent overly long speech
                    MAX_TTS_TEXT_LENGTH = 600
                    tts_text = response_text.replace('*', '')
                    if len(tts_text) > MAX_TTS_TEXT_LENGTH:
                        tts_text = tts_text[:MAX_TTS_TEXT_LENGTH].rsplit(' ', 1)[0]
                    
                    tts_payload = {
                        "text": tts_text,
                        "speaker": tts_speaker
                    }
                    if tts_model:
                        tts_payload["model_name"] = tts_model
                    if tts_engine:
                        tts_payload["engine"] = tts_engine
                    
                    tts_response = requests.post(
                        tts_url,
                        json=tts_payload,
                        timeout=tts_timeout
                    )
                    if tts_response.status_code == 200:
                        broadcast_sent = True
                        print("DEBUG: TTS broadcast successful")
                    else:
                        print(f"DEBUG: TTS broadcast failed with status {tts_response.status_code}")
                        print(f"DEBUG: TTS response: {tts_response.text}")
                except Exception as e:
                    print(f"ERROR: TTS broadcast failed: {str(e)}")
            elif broadcast and not tts_speaker:
                print("DEBUG: TTS broadcast requested but speaker not specified")
            else:
                print("DEBUG: TTS broadcast requested but TTS_BROADCAST_URL not configured")
        
        # === RETURN RESPONSE ===
        return jsonify({
            "success": True,
            "query": transcribed_text,
            "response": response_text,
            "tools_used": tools_used,
            "broadcast_sent": broadcast_sent
        })
    
    except Exception as e:
        print(f"ERROR: Voice query failed: {str(e)}")
        import traceback
        print(f"DEBUG: Full traceback: {traceback.format_exc()}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500